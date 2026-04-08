import os
import re

import streamlit as st
from docx import Document
import numpy as np

# CONFIG
DOCS_FOLDER = "docs"
EMBED_MODEL_NAME = "all-MiniLM-L6-v2"
TOP_K = 3

# Reduce noisy tokenizer threading warnings in constrained local environments.
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

# LOAD DOCX
def extract_docx_content(file_path: str) -> str:
    """
    Extract paragraphs + tables from a .docx file.
    Screenshots/images OCR is skipped here for simplicity.
    """
    doc = Document(file_path)
    parts = []

    # Paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)

# CHUNKING
def chunk_text(text: str, chunk_size: int = 700, overlap: int = 120):
    """
    Basic character-based chunking.
    """
    text = re.sub(r"\s+", " ", text).strip()
    chunks = []

    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start += (chunk_size - overlap)

    return chunks

# BUILD VECTOR INDEX
@st.cache_resource
def build_knowledge_base():
    from sentence_transformers import SentenceTransformer

    documents = []
    metadata = []

    if not os.path.isdir(DOCS_FOLDER):
        raise ValueError(f"Missing docs folder: {DOCS_FOLDER}")

    for filename in os.listdir(DOCS_FOLDER):
        if filename.endswith(".docx"):
            path = os.path.join(DOCS_FOLDER, filename)
            content = extract_docx_content(path)
            chunks = chunk_text(content)

            for i, chunk in enumerate(chunks):
                documents.append(chunk)
                metadata.append({
                    "source": filename,
                    "chunk_id": i
                })

    if not documents:
        raise ValueError("No .docx documents found in docs/ folder.")

    model = SentenceTransformer(EMBED_MODEL_NAME)
    embeddings = model.encode(documents, convert_to_numpy=True).astype("float32")

    return model, embeddings, documents, metadata

# RETRIEVAL
def retrieve(query: str, model, embeddings, documents, metadata, top_k: int = TOP_K):
    
    query_emb = model.encode([query], convert_to_numpy=True).astype("float32")[0]
    distances = np.linalg.norm(embeddings - query_emb, axis=1)
    indices = np.argsort(distances)[:top_k]

    results = []
    for idx in indices:
        results.append({
            "text": documents[idx],
            "source": metadata[idx]["source"],
            "chunk_id": metadata[idx]["chunk_id"],
            "distance": float(distances[idx])
        })
    return results

# PROMPT BUILDING
def build_prompt(user_query: str, retrieved_chunks):
    context = "\n\n".join(
        [f"[Source: {c['source']} | Chunk {c['chunk_id']}]\n{c['text']}" for c in retrieved_chunks]
    )

    prompt = f"""
You are an internal incident-support assistant.

Your role:
- Help the user understand possible resolution steps for a ServiceNow incident
- Use ONLY the provided context
- If unsure, say "Not enough information in knowledge base"
- Be practical and concise
- Output:
  1. Incident type guess
  2. Suggested resolution steps
  3. SQL/action examples if found
  4. Source references used

Context:
{context}

User incident:
{user_query}
"""
    return prompt.strip()

# LOCAL LLM CALL (OLLAMA)
def ask_ollama(prompt: str, model_name: str = "llama3:8b"):
    """
    Requires local Ollama running:
    - ollama serve
    - ollama pull llama3
    """
    try:
        import ollama
        response = ollama.chat(
            model=model_name,
            messages=[{"role": "user", "content": prompt}]
        )
        return response["message"]["content"]
    except Exception as e:
        return f"Local LLM error: {e}"

# UI
def main():
    st.set_page_config(page_title="Incident Chatbot", layout="wide")
    st.title("ServiceNow Incident Analysis Chatbot")
    st.caption("Local RAG prototype using .docx knowledge base")

    st.info("Knowledge is loaded from .docx files in the docs folder.")

    try:
        model, embeddings, documents, metadata = build_knowledge_base()
    except Exception as e:
        st.error(f"Failed to build knowledge base: {e}")
        st.stop()

    user_query = st.text_area(
        "Paste incident description:",
        height=180,
        placeholder="Example: Please delete duplicate yellow-marked records from NRI..."
    )

    if st.button("Analyze Incident"):
        if not user_query.strip():
            st.warning("Please enter an incident description.")
            return

        retrieved = retrieve(user_query, model, embeddings, documents, metadata, top_k=TOP_K)
        prompt = build_prompt(user_query, retrieved)
        answer = ask_ollama(prompt, model_name="llama3:8b")

        st.subheader("Recommended Answer")
        st.write(answer)

        st.subheader("Retrieved Knowledge")
        for r in retrieved:
            with st.expander(f"{r['source']} | chunk {r['chunk_id']}"):
                st.write(r["text"])

if __name__ == "__main__":
    main()